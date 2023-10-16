import streamlit as st
import os
import requests
from docx import Document
from docx.shared import Inches
from io import BytesIO
import docx2pdf
from PIL import Image
from transformers import AutoModelForCausalLM, GenerationConfig

import torch
import random

if os.name == 'nt':
    import pythoncom
    pythoncom.CoInitialize()
if os.name == 'posix':
    import subprocess
    def convert_docx_to_pdf(docx_path, pdf_path):
        subprocess.run(['unoconv', '-f', 'pdf', '-o', pdf_path, docx_path])

def get_image_data(url):
    response = requests.get(url)
    response.raise_for_status()
    return response.content

if 'chapter_count' not in st.session_state:
    st.session_state.chapter_count = 1
if 'file_available' not in st.session_state:
    st.session_state.file_available = False
if 'images' not in st.session_state:
    st.session_state.images = []
if 'cur_time' not in st.session_state:
    st.session_state.cur_time = 0
if 'screenplays' not in st.session_state:
    st.session_state.screenplays = []
if 'subtitles' not in st.session_state:
    st.session_state.subtitles = []
# initialize document format
screenplay_document = Document()
section = screenplay_document.sections[0]
section.page_width = Inches(8.27)  # Width for A4 paper
section.page_height = Inches(11.69)  # Height for A4 paper

section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
screenplay_document = Document()

subtitle_document = Document()
section = subtitle_document.sections[0]
section.page_width = Inches(8.27)  # Width for A4 paper
section.page_height = Inches(11.69)  # Height for A4 paper

section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)


st.write('Film Title')
content_text = st.text_input('Write your contents of your film')
chapter_count = st.text_input('Write the count of parts')
generate_btn = st.button('Generate screenplay and subtitles')

if content_text and generate_btn and chapter_count:
    st.session_state.file_available = False
    st.session_state.chapter_count = int(chapter_count)
    st.session_state.cur_time = random.randint(1000, 3000)
    st.session_state.screenplays = []
    st.session_state.subtitles = []

    first_prompt = f'write the first part of screenplay of film about {content_text}. write as much as possible.'
    general_prompt = 'write the next part as much as possible'
    last_prompt = 'write the next part as much as possible and end the screenplay'
    
    model = AutoModelForCausalLM.from_pretrained("model/txt2txt_fine_tuned_005")
    generation_config = GenerationConfig(max_new_tokens=4096, do_sample=True, top_k=50, eos_token_id=model.config.eos_token_id)
    # generate screenplays
    chapter_count = st.session_state.chapter_count
    for i in range(chapter_count):
        prompt = ''
        if i == 0: 
            prompt = first_prompt
            response = model.generate(**prompt, generation_config=generation_config)
        elif i == chapter_count - 1:    
            prompt = last_prompt
            last_text = st.session_state.screenplays[len(st.session_state.screenplays) - 1]

            response = model.generate(**prompt, generation_config=generation_config)
        else:
            prompt = general_prompt
            last_text = st.session_state.screenplays[len(st.session_state.screenplays) - 1]

            response = model.generate(**prompt, generation_config=generation_config)

        screenplay = response['choices'][0]['message']['content']
        st.session_state.screenplays.append(screenplay)

        screenplay_expander = st.expander(f'Screenplay Part {i + 1}')
        screenplay_expander.text(screenplay)
        
        extract_prompt = f'''"{screenplay}"\nextract dialogues from here and write only dialogues (no name) like this.\n"[dialogue1]\n[dialouge2]\n[dialogue3]\n..."'''
        
        subtitles = response['choices'][0]['message']['content']
        subtitle_texts = ""
        for subtitle in subtitles.split('\n'):
            length = len(subtitle)
            talk_time = random.randint(65 * length, 75 * length)
            space_time = random.randint(500, 1000)
            start_time = st.session_state.cur_time
            end_time = st.session_state.cur_time + talk_time
            _h, _m, _s, _ms = start_time // 1000 // 3600, start_time // 1000 % 3600 // 60, start_time // 1000 % 60, start_time % 1000
            start_text = f"{_h:02d}:{_m:02d}:{_s:02d},{_ms:03d}"
            _h, _m, _s, _ms = end_time // 1000 // 3600, end_time // 1000 % 3600 // 60, end_time // 1000 % 60, end_time % 1000
            end_text = f"{_h:02d}:{_m:02d}:{_s:02d},{_ms:03d}"
            
            subtitle_text = f"{start_text} --> {end_text}\n{subtitle}\n\n"
            subtitle_texts = subtitle_texts + subtitle_text
            print (subtitle_text)
            st.session_state.cur_time = st.session_state.cur_time + talk_time + space_time
        
        st.session_state.subtitles.append(subtitle_texts)

        subtitle_expander = st.expander(f'Subtitle Part {i + 1}')
        subtitle_expander.text(subtitle_texts)

        screenplay_document.add_paragraph(screenplay)
        subtitle_document.add_paragraph(subtitle_texts)

    screenplay_document.save('screenplay.docx')
    subtitle_document.save('subtitle.docx')
    if os.name == 'nt':
        docx2pdf.convert('screenplay.docx', 'screenplay.pdf')
        docx2pdf.convert('subtitle.docx', 'subtitle.pdf')
    if os.name == 'posix':
        convert_docx_to_pdf('screenplay.docx', 'screenplay.pdf')
        convert_docx_to_pdf('subtitle.docx', 'subtitle.pdf')
    st.session_state.file_available = True
    st.experimental_rerun()

if st.session_state.file_available:
    for i in range(st.session_state.chapter_count):
        screenplay_expander = st.expander(f'Screenplay Part {i + 1}')
        screenplay_expander.text(st.session_state.screenplays[i])

        subtitle_expander = st.expander(f'Subtitle Part {i + 1}')
        subtitle_expander.text(st.session_state.subtitles[i])
        # expander.image(st.session_state.images[i])
    col1, col2, col3, col4, col5 = st.columns(5)
    with open("screenplay.docx", "rb") as file:
        col2.download_button(
                label="Download screenplay docx",
                data=file,
                file_name="screenplay.docx"
            )
    with open("screenplay.pdf", "rb") as file:
        col4.download_button(
                label="Download screenplay pdf",
                data=file,
                file_name="screenplay.pdf"
            )
    col1, col2, col3, col4, col5 = st.columns(5)
        
    with open("subtitle.docx", "rb") as file:
        col2.download_button(
                label="Download subtitle docx",
                data=file,
                file_name="subtitle.docx"
            )
    with open("subtitle.pdf", "rb") as file:
        col4.download_button(
                label="Download subtitle pdf",
                data=file,
                file_name="subtitle.pdf"
            )